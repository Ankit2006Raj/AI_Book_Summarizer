import os
import json
import csv
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import Config

class Exporter:
    """Export summaries and data in various formats"""
    
    def __init__(self):
        self.export_dir = Config.EXPORT_DIR
    
    def export_to_pdf(self, summary_data, filename='book_summary.pdf'):
        """
        Export summary to PDF
        
        Args:
            summary_data (dict): Summary data
            filename (str): Output filename
            
        Returns:
            str: Path to exported file
        """
        output_path = os.path.join(self.export_dir, filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#34495e'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        normal_style = styles['Normal']
        
        # Cover Page
        book_info = summary_data.get('book_info', {})
        elements.append(Spacer(1, 2*inch))
        elements.append(Paragraph(f"<b>{book_info.get('title', 'Untitled')}</b>", title_style))
        elements.append(Spacer(1, 0.3*inch))
        elements.append(Paragraph(f"by {book_info.get('author', 'Unknown')}", styles['Heading3']))
        elements.append(Spacer(1, 0.5*inch))
        elements.append(Paragraph("<b>AI-Generated Summary</b>", styles['Heading3']))
        elements.append(Spacer(1, 0.2*inch))
        elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", normal_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Book statistics
        stats_data = [
            ['Total Chapters:', str(book_info.get('total_chapters', 0))],
            ['Total Words:', f"{book_info.get('total_words', 0):,}"],
            ['Total Pages:', str(book_info.get('total_pages', 0))],
            ['Est. Reading Time:', f"{book_info.get('estimated_reading_time', 0):.1f} minutes"],
        ]
        
        stats_table = Table(stats_data, colWidths=[2.5*inch, 2*inch])
        stats_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        elements.append(stats_table)
        elements.append(Spacer(1, 0.3*inch))
        elements.append(Paragraph("<i>Powered by Google Gemini AI</i>", styles['Italic']))
        elements.append(PageBreak())
        
        # Overall Summary
        elements.append(Paragraph("Overall Summary", heading_style))
        elements.append(Spacer(1, 0.1*inch))
        overall_summary = summary_data.get('overall_summary', '')
        for para in overall_summary.split('\n\n'):
            if para.strip():
                elements.append(Paragraph(para.strip(), normal_style))
                elements.append(Spacer(1, 0.1*inch))
        
        elements.append(PageBreak())
        
        # Themes and Analysis
        themes = summary_data.get('overall_themes', {})
        if themes:
            elements.append(Paragraph("Themes & Analysis", heading_style))
            elements.append(Spacer(1, 0.1*inch))
            
            if themes.get('genre'):
                elements.append(Paragraph(f"<b>Genre:</b> {themes['genre']}", normal_style))
                elements.append(Spacer(1, 0.05*inch))
            
            if themes.get('tone'):
                elements.append(Paragraph(f"<b>Tone:</b> {themes['tone']}", normal_style))
                elements.append(Spacer(1, 0.05*inch))
            
            if themes.get('themes'):
                themes_text = ', '.join(themes['themes'])
                elements.append(Paragraph(f"<b>Main Themes:</b> {themes_text}", normal_style))
                elements.append(Spacer(1, 0.1*inch))
            
            if themes.get('emotional_journey'):
                elements.append(Paragraph(f"<b>Emotional Journey:</b> {themes['emotional_journey']}", normal_style))
            
            elements.append(PageBreak())
        
        # Chapter Summaries
        elements.append(Paragraph("Chapter-by-Chapter Summary", heading_style))
        elements.append(Spacer(1, 0.2*inch))
        
        for chapter in summary_data.get('chapters', []):
            chapter_title = chapter.get('chapter_title', 'Untitled')
            chapter_num = chapter.get('chapter_number', '')
            
            elements.append(Paragraph(f"<b>Chapter {chapter_num}: {chapter_title}</b>", styles['Heading3']))
            elements.append(Spacer(1, 0.1*inch))
            
            summary = chapter.get('summary', '')
            for para in summary.split('\n\n'):
                if para.strip():
                    elements.append(Paragraph(para.strip(), normal_style))
                    elements.append(Spacer(1, 0.1*inch))
            
            # Key points
            key_concepts = chapter.get('key_concepts', [])
            if key_concepts:
                elements.append(Paragraph("<b>Key Points:</b>", normal_style))
                for concept in key_concepts[:5]:
                    elements.append(Paragraph(f"• {concept}", normal_style))
                elements.append(Spacer(1, 0.1*inch))
            
            elements.append(Spacer(1, 0.2*inch))
        
        # Build PDF
        doc.build(elements)
        
        return output_path
    
    def export_to_docx(self, summary_data, filename='book_summary.docx'):
        """
        Export summary to DOCX
        
        Args:
            summary_data (dict): Summary data
            filename (str): Output filename
            
        Returns:
            str: Path to exported file
        """
        output_path = os.path.join(self.export_dir, filename)
        
        doc = Document()
        
        # Cover Page
        book_info = summary_data.get('book_info', {})
        
        title = doc.add_heading(book_info.get('title', 'Untitled'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        author = doc.add_paragraph(f"by {book_info.get('author', 'Unknown')}")
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author.runs[0].font.size = Pt(14)
        
        doc.add_paragraph()
        
        subtitle = doc.add_paragraph('AI-Generated Summary')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].bold = True
        
        date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Statistics table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        stats = [
            ('Total Chapters:', str(book_info.get('total_chapters', 0))),
            ('Total Words:', f"{book_info.get('total_words', 0):,}"),
            ('Total Pages:', str(book_info.get('total_pages', 0))),
            ('Est. Reading Time:', f"{book_info.get('estimated_reading_time', 0):.1f} minutes"),
        ]
        
        for idx, (label, value) in enumerate(stats):
            row = table.rows[idx]
            row.cells[0].text = label
            row.cells[1].text = value
        
        doc.add_paragraph()
        
        footer = doc.add_paragraph('Powered by Google Gemini AI')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].italic = True
        footer.runs[0].font.size = Pt(10)
        
        doc.add_page_break()
        
        # Overall Summary
        doc.add_heading('Overall Summary', 1)
        overall_summary = summary_data.get('overall_summary', '')
        doc.add_paragraph(overall_summary)
        
        doc.add_page_break()
        
        # Themes
        themes = summary_data.get('overall_themes', {})
        if themes:
            doc.add_heading('Themes & Analysis', 1)
            
            if themes.get('genre'):
                p = doc.add_paragraph()
                p.add_run('Genre: ').bold = True
                p.add_run(themes['genre'])
            
            if themes.get('tone'):
                p = doc.add_paragraph()
                p.add_run('Tone: ').bold = True
                p.add_run(themes['tone'])
            
            if themes.get('themes'):
                p = doc.add_paragraph()
                p.add_run('Main Themes: ').bold = True
                p.add_run(', '.join(themes['themes']))
            
            if themes.get('emotional_journey'):
                p = doc.add_paragraph()
                p.add_run('Emotional Journey: ').bold = True
                p.add_run(themes['emotional_journey'])
            
            doc.add_page_break()
        
        # Chapter Summaries
        doc.add_heading('Chapter-by-Chapter Summary', 1)
        
        for chapter in summary_data.get('chapters', []):
            chapter_title = chapter.get('chapter_title', 'Untitled')
            chapter_num = chapter.get('chapter_number', '')
            
            doc.add_heading(f"Chapter {chapter_num}: {chapter_title}", 2)
            
            summary = chapter.get('summary', '')
            doc.add_paragraph(summary)
            
            # Key points
            key_concepts = chapter.get('key_concepts', [])
            if key_concepts:
                p = doc.add_paragraph()
                p.add_run('Key Points:').bold = True
                for concept in key_concepts[:5]:
                    doc.add_paragraph(concept, style='List Bullet')
        
        doc.save(output_path)
        
        return output_path
    
    def export_to_txt(self, summary_data, filename='book_summary.txt'):
        """
        Export summary to plain text
        
        Args:
            summary_data (dict): Summary data
            filename (str): Output filename
            
        Returns:
            str: Path to exported file
        """
        output_path = os.path.join(self.export_dir, filename)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            book_info = summary_data.get('book_info', {})
            
            # Header
            f.write("=" * 80 + "\n")
            f.write(f"{book_info.get('title', 'Untitled').upper()}\n")
            f.write(f"by {book_info.get('author', 'Unknown')}\n")
            f.write("=" * 80 + "\n\n")
            
            f.write("AI-GENERATED SUMMARY\n")
            f.write(f"Generated on: {datetime.now().strftime('%B %d, %Y')}\n\n")
            
            # Statistics
            f.write(f"Total Chapters: {book_info.get('total_chapters', 0)}\n")
            f.write(f"Total Words: {book_info.get('total_words', 0):,}\n")
            f.write(f"Total Pages: {book_info.get('total_pages', 0)}\n")
            f.write(f"Est. Reading Time: {book_info.get('estimated_reading_time', 0):.1f} minutes\n\n")
            
            f.write("Powered by Google Gemini AI\n")
            f.write("=" * 80 + "\n\n")
            
            # Overall Summary
            f.write("OVERALL SUMMARY\n")
            f.write("-" * 80 + "\n")
            f.write(summary_data.get('overall_summary', '') + "\n\n")
            
            # Themes
            themes = summary_data.get('overall_themes', {})
            if themes:
                f.write("THEMES & ANALYSIS\n")
                f.write("-" * 80 + "\n")
                if themes.get('genre'):
                    f.write(f"Genre: {themes['genre']}\n")
                if themes.get('tone'):
                    f.write(f"Tone: {themes['tone']}\n")
                if themes.get('themes'):
                    f.write(f"Main Themes: {', '.join(themes['themes'])}\n")
                if themes.get('emotional_journey'):
                    f.write(f"Emotional Journey: {themes['emotional_journey']}\n")
                f.write("\n")
            
            # Chapter Summaries
            f.write("CHAPTER-BY-CHAPTER SUMMARY\n")
            f.write("-" * 80 + "\n\n")
            
            for chapter in summary_data.get('chapters', []):
                chapter_title = chapter.get('chapter_title', 'Untitled')
                chapter_num = chapter.get('chapter_number', '')
                
                f.write(f"CHAPTER {chapter_num}: {chapter_title.upper()}\n")
                f.write("-" * 40 + "\n")
                f.write(chapter.get('summary', '') + "\n\n")
                
                # Key points
                key_concepts = chapter.get('key_concepts', [])
                if key_concepts:
                    f.write("Key Points:\n")
                    for concept in key_concepts[:5]:
                        f.write(f"  • {concept}\n")
                    f.write("\n")
        
        return output_path
    
    def export_to_json(self, summary_data, filename='book_summary.json'):
        """
        Export summary to JSON
        
        Args:
            summary_data (dict): Summary data
            filename (str): Output filename
            
        Returns:
            str: Path to exported file
        """
        output_path = os.path.join(self.export_dir, filename)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(summary_data, f, ensure_ascii=False, indent=2)
        
        return output_path
    
    def export_quiz_to_csv(self, quiz_data, filename='quiz.csv'):
        """
        Export quiz to CSV
        
        Args:
            quiz_data (list): Quiz questions
            filename (str): Output filename
            
        Returns:
            str: Path to exported file
        """
        output_path = os.path.join(self.export_dir, filename)
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Type', 'Question', 'Options', 'Correct Answer', 'Explanation'])
            
            for question in quiz_data:
                options = json.dumps(question.get('options', []))
                writer.writerow([
                    question.get('type', ''),
                    question.get('question', ''),
                    options,
                    question.get('correct_answer', ''),
                    question.get('explanation', '')
                ])
        
        return output_path
    
    def export_flashcards(self, summary_data, filename='flashcards.txt'):
        """
        Export flashcards for Anki/Notion
        
        Args:
            summary_data (dict): Summary data
            filename (str): Output filename
            
        Returns:
            str: Path to exported file
        """
        output_path = os.path.join(self.export_dir, filename)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for chapter in summary_data.get('chapters', []):
                chapter_title = chapter.get('chapter_title', 'Untitled')
                
                # Key concepts as flashcards
                for concept in chapter.get('key_concepts', []):
                    f.write(f"Q: What is the key concept in {chapter_title}?\n")
                    f.write(f"A: {concept}\n\n")
                
                # Keywords as flashcards
                for keyword in chapter.get('keywords', [])[:3]:
                    f.write(f"Q: Define '{keyword}' in the context of {chapter_title}\n")
                    f.write(f"A: [Student should provide definition based on chapter]\n\n")
        
        return output_path
